"""
Document Processing Pipeline — Read any legal document the user uploads.
========================================================================

Handles:
- PDFs (digital text + scanned images)
- Images (JPG, PNG - scanned FIRs, old court orders, evidence photos)
- Word documents (.docx)
- Plain text files

Features:
- PDF text extraction (pdfplumber for digital, OCR for scanned)
- Multi-language OCR (English + Hindi + regional via EasyOCR)
- Evidence photo analysis (sends to LLM for visual description)
- Section/citation extraction from parsed text
- Structured output: text, metadata, extracted entities
"""
from __future__ import annotations

import base64
import io
import logging
import mimetypes
import os
import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from PIL import Image

logger = logging.getLogger(__name__)

# Lazy-loaded heavy modules
_pdfplumber = None
_easyocr_reader = None
_docx = None


def _get_pdfplumber():
    global _pdfplumber
    if _pdfplumber is None:
        import pdfplumber
        _pdfplumber = pdfplumber
    return _pdfplumber


def _get_easyocr(languages: List[str] = None):
    """Get or create EasyOCR reader. Lazy-loaded and cached."""
    global _easyocr_reader
    if _easyocr_reader is None:
        import easyocr
        langs = languages or ["en", "hi"]
        logger.info(f"Initializing EasyOCR with languages: {langs}")
        _easyocr_reader = easyocr.Reader(langs, gpu=True)
        logger.info("EasyOCR ready (GPU-accelerated)")
    return _easyocr_reader


def _get_docx():
    global _docx
    if _docx is None:
        import docx
        _docx = docx
    return _docx


# ── Data Classes ──────────────────────────────────────────────────────────

@dataclass
class ExtractedDocument:
    """Result of document processing."""
    filename: str
    file_type: str  # pdf, image, docx, txt
    text: str  # Full extracted text
    page_count: int = 1
    language_detected: str = "en"
    sections_mentioned: List[str] = field(default_factory=list)
    citations_found: List[str] = field(default_factory=list)
    parties_mentioned: List[str] = field(default_factory=list)
    dates_found: List[str] = field(default_factory=list)
    doc_type_guess: str = ""  # FIR, charge_sheet, court_order, judgment, evidence, petition, etc.
    confidence: float = 1.0
    evidence_description: str = ""  # For photos: what the image shows
    metadata: Dict[str, Any] = field(default_factory=dict)

    def summary(self, max_len: int = 2000) -> str:
        """Get a summary suitable for LLM context."""
        parts = [f"[Document: {self.filename} | Type: {self.doc_type_guess or self.file_type}]"]
        if self.evidence_description:
            parts.append(f"[Visual: {self.evidence_description}]")
        parts.append(self.text[:max_len])
        if self.sections_mentioned:
            parts.append(f"[Sections cited: {', '.join(self.sections_mentioned[:15])}]")
        if self.citations_found:
            parts.append(f"[Case citations: {', '.join(self.citations_found[:10])}]")
        return "\n".join(parts)


# ── Document Processor ────────────────────────────────────────────────────

class DocumentProcessor:
    """
    Processes uploaded legal documents into structured text.
    
    Uses:
    - pdfplumber for digital PDFs
    - EasyOCR (GPU) for scanned documents and images (multi-language)
    - LLM for evidence photo description
    - Regex for legal entity extraction
    """

    def __init__(self, llm_client=None, ocr_languages: List[str] = None):
        """
        Args:
            llm_client: Optional LLM client for evidence photo analysis.
            ocr_languages: OCR language codes. Default: ["en", "hi"]
        """
        self.llm_client = llm_client
        self.ocr_languages = ocr_languages or ["en", "hi"]

    def process_file(self, file_path: str = None, file_bytes: bytes = None,
                     filename: str = "document", content_type: str = None) -> ExtractedDocument:
        """
        Process a single document file.
        
        Args:
            file_path: Path to file on disk.
            file_bytes: Raw file bytes (for uploads).
            filename: Original filename.
            content_type: MIME type hint.
        
        Returns:
            ExtractedDocument with extracted text and metadata.
        """
        if file_path and not file_bytes:
            with open(file_path, "rb") as f:
                file_bytes = f.read()
            filename = filename or os.path.basename(file_path)

        if not file_bytes:
            raise ValueError("Either file_path or file_bytes must be provided")

        # Detect file type
        ext = Path(filename).suffix.lower()
        if not content_type:
            content_type = mimetypes.guess_type(filename)[0] or ""

        logger.info(f"Processing document: {filename} ({ext}, {len(file_bytes)} bytes)")

        if ext == ".pdf" or "pdf" in content_type:
            return self._process_pdf(file_bytes, filename)
        elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp") or "image" in content_type:
            return self._process_image(file_bytes, filename)
        elif ext == ".docx" or "wordprocessing" in content_type:
            return self._process_docx(file_bytes, filename)
        elif ext in (".txt", ".text") or "text/plain" in content_type:
            return self._process_text(file_bytes, filename)
        else:
            # Try as text first, fall back to image OCR
            try:
                text = file_bytes.decode("utf-8", errors="strict")
                return self._process_text(text.encode(), filename)
            except (UnicodeDecodeError, ValueError):
                return self._process_image(file_bytes, filename)

    def process_multiple(self, files: List[Dict[str, Any]]) -> List[ExtractedDocument]:
        """
        Process multiple documents.
        
        Args:
            files: List of dicts with keys: file_bytes, filename, content_type
        
        Returns:
            List of ExtractedDocument results.
        """
        results = []
        for f in files:
            try:
                doc = self.process_file(
                    file_bytes=f.get("file_bytes"),
                    filename=f.get("filename", "document"),
                    content_type=f.get("content_type"),
                )
                results.append(doc)
            except Exception as e:
                logger.error(f"Failed to process {f.get('filename', '?')}: {e}")
                results.append(ExtractedDocument(
                    filename=f.get("filename", "unknown"),
                    file_type="error",
                    text=f"[Error processing document: {str(e)}]",
                    confidence=0.0,
                ))
        return results

    # ── PDF Processing ────────────────────────────────────────────────────

    def _process_pdf(self, data: bytes, filename: str) -> ExtractedDocument:
        """Extract text from PDF — tries digital extraction first, falls back to OCR."""
        pdfplumber = _get_pdfplumber()

        pages_text = []
        page_count = 0
        needs_ocr = False

        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    if text.strip():
                        pages_text.append(text)
                    else:
                        needs_ocr = True
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}. Falling back to OCR.")
            needs_ocr = True

        # If most pages had no text, it's a scanned PDF — use OCR
        digital_text = "\n\n".join(pages_text)
        if not digital_text.strip() or needs_ocr or len(digital_text) < 100:
            logger.info(f"PDF appears scanned or has images. Running OCR on {page_count} pages...")
            ocr_text = self._ocr_pdf_pages(data)
            if ocr_text.strip():
                full_text = digital_text + "\n\n" + ocr_text if digital_text.strip() else ocr_text
            else:
                full_text = digital_text
        else:
            full_text = digital_text

        doc = ExtractedDocument(
            filename=filename,
            file_type="pdf",
            text=full_text.strip(),
            page_count=page_count,
        )
        self._extract_legal_entities(doc)
        self._guess_document_type(doc)
        return doc

    def _ocr_pdf_pages(self, pdf_bytes: bytes) -> str:
        """Convert PDF pages to images and run OCR."""
        try:
            # Convert PDF pages to images using pdfplumber
            pdfplumber = _get_pdfplumber()
            reader = _get_easyocr(self.ocr_languages)
            texts = []

            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for i, page in enumerate(pdf.pages):
                    try:
                        img = page.to_image(resolution=200)
                        # Save to bytes
                        buf = io.BytesIO()
                        img.save(buf, format="PNG")
                        buf.seek(0)
                        img_bytes = buf.read()

                        results = reader.readtext(img_bytes)
                        page_text = " ".join([r[1] for r in results])
                        if page_text.strip():
                            texts.append(f"[Page {i+1}]\n{page_text}")
                    except Exception as e:
                        logger.warning(f"OCR failed on page {i+1}: {e}")

            return "\n\n".join(texts)
        except Exception as e:
            logger.error(f"PDF OCR failed: {e}")
            return ""

    # ── Image Processing ──────────────────────────────────────────────────

    def _process_image(self, data: bytes, filename: str) -> ExtractedDocument:
        """Process an image — OCR for text, LLM for evidence description."""
        # Run OCR
        ocr_text = self._ocr_image(data)

        # If it looks like evidence (not a document), also get visual description
        evidence_desc = ""
        is_evidence = self._is_evidence_photo(ocr_text, filename)

        if is_evidence and self.llm_client:
            evidence_desc = self._describe_evidence(data, filename)

        doc = ExtractedDocument(
            filename=filename,
            file_type="image",
            text=ocr_text.strip() if ocr_text.strip() else "[Image with no extractable text]",
            evidence_description=evidence_desc,
        )

        if ocr_text.strip():
            self._extract_legal_entities(doc)
            self._guess_document_type(doc)
        elif is_evidence:
            doc.doc_type_guess = "evidence_photo"

        return doc

    def _ocr_image(self, data: bytes) -> str:
        """Run EasyOCR on an image."""
        try:
            reader = _get_easyocr(self.ocr_languages)
            results = reader.readtext(data)
            # Sort by position (top-to-bottom, left-to-right)
            results.sort(key=lambda r: (r[0][0][1], r[0][0][0]))
            text = "\n".join([r[1] for r in results if r[2] > 0.3])  # confidence > 0.3
            return text
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return ""

    def _is_evidence_photo(self, ocr_text: str, filename: str) -> bool:
        """Heuristic: is this a photo of physical evidence vs. a scanned document?"""
        fn = filename.lower()
        if any(w in fn for w in ["evidence", "photo", "scene", "weapon", "injury", "exhibit"]):
            return True
        # Low text content suggests it's a photo, not a document
        if len(ocr_text.strip()) < 50:
            return True
        return False

    def _describe_evidence(self, data: bytes, filename: str) -> str:
        """Use LLM (GPT-4o vision) to describe evidence in a photo."""
        if not self.llm_client:
            return ""
        try:
            b64_image = base64.b64encode(data).decode("utf-8")

            # Determine image type
            ext = Path(filename).suffix.lower()
            mime = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "webp": "webp"}.get(ext.lstrip("."), "jpeg")

            messages = [
                {"role": "system", "content": (
                    "You are a forensic evidence analyst assisting a legal case. "
                    "Describe what you see in this image objectively and thoroughly. "
                    "Note: physical condition of objects, any visible marks, injuries, "
                    "weapons, documents, locations, people (without identifying). "
                    "Be factual and precise. This description will be used in legal proceedings."
                )},
                {"role": "user", "content": [
                    {"type": "text", "text": f"Describe this evidence photo ({filename}):"},
                    {"type": "image_url", "image_url": {"url": f"data:image/{mime};base64,{b64_image}"}},
                ]},
            ]
            response = self.llm_client.chat(messages, temperature=0.2, max_tokens=500)
            return response.strip()
        except Exception as e:
            logger.warning(f"Evidence description failed: {e}")
            return ""

    # ── DOCX Processing ───────────────────────────────────────────────────

    def _process_docx(self, data: bytes, filename: str) -> ExtractedDocument:
        """Extract text from Word documents."""
        docx = _get_docx()
        doc_obj = docx.Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc_obj.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        doc = ExtractedDocument(
            filename=filename,
            file_type="docx",
            text=full_text.strip(),
            page_count=max(1, len(full_text) // 3000),  # Estimate
        )
        self._extract_legal_entities(doc)
        self._guess_document_type(doc)
        return doc

    # ── Text Processing ───────────────────────────────────────────────────

    def _process_text(self, data: bytes, filename: str) -> ExtractedDocument:
        """Process plain text files."""
        for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
            try:
                text = data.decode(encoding)
                break
            except (UnicodeDecodeError, ValueError):
                continue
        else:
            text = data.decode("utf-8", errors="replace")

        doc = ExtractedDocument(
            filename=filename,
            file_type="txt",
            text=text.strip(),
        )
        self._extract_legal_entities(doc)
        self._guess_document_type(doc)
        return doc

    # ── Legal Entity Extraction ───────────────────────────────────────────

    def _extract_legal_entities(self, doc: ExtractedDocument) -> None:
        """Extract legal sections, citations, parties, dates from text."""
        text = doc.text

        # Sections (IPC, BNS, CrPC, BNSS, CPC, etc.)
        section_patterns = [
            r"[Ss]ection\s+(\d+[A-Z]?(?:/\d+)?)\s+(?:of\s+)?(?:the\s+)?(?:IPC|Indian\s+Penal\s+Code)",
            r"[Ss]ection\s+(\d+[A-Z]?)\s+(?:of\s+)?(?:the\s+)?(?:BNS|Bharatiya\s+Nyaya)",
            r"[Ss]ection\s+(\d+[A-Z]?)\s+(?:of\s+)?(?:the\s+)?(?:Cr\.?P\.?C|CrPC|Code\s+of\s+Criminal)",
            r"[Ss]ection\s+(\d+[A-Z]?)\s+(?:of\s+)?(?:the\s+)?(?:BNSS|Bharatiya\s+Nagarik)",
            r"[Ss]ection\s+(\d+[A-Z]?)\s+(?:of\s+)?(?:the\s+)?(?:CPC|C\.?P\.?C|Code\s+of\s+Civil)",
            r"[Ss]ection\s+(\d+[A-Z]?)\s+(?:of\s+)?(?:the\s+)?(?:Evidence|BSA)",
            r"[Aa]rticle\s+(\d+[A-Z]?)\s+(?:of\s+)?(?:the\s+)?Constitution",
            r"[Ss]ection\s+(\d+[A-Z]?)\s+(?:of\s+)?(?:the\s+)?NDPS",
            r"[Ss]ection\s+(\d+[A-Z]?)\s+(?:of\s+)?(?:the\s+)?POCSO",
            r"[Ss]ection\s+(\d+[A-Z]?)\s+(?:of\s+)?(?:the\s+)?(?:NI|Negotiable\s+Instruments)",
            r"[Ss]ection\s+(\d+[A-Z]?)\s+(?:of\s+)?(?:the\s+)?(?:DV|Domestic\s+Violence)",
            r"[Ss]ection\s+(\d+[A-Z]?)\s+(?:of\s+)?(?:the\s+)?(?:IT\s+Act|Information\s+Technology)",
        ]
        sections = set()
        for p in section_patterns:
            for m in re.finditer(p, text, re.IGNORECASE):
                sections.add(m.group(0).strip())
        doc.sections_mentioned = sorted(sections)[:30]

        # Case citations
        citation_patterns = [
            r"\(\d{4}\)\s+\d+\s+SCC\s+\d+",
            r"AIR\s+\d{4}\s+\w+\s+\d+",
            r"\d{4}\s+\(\d+\)\s+SCC\s+\d+",
            r"\d{4}\s+Cri\.?L\.?J\.?\s+\d+",
            r"\d{4}\s+SCC\s+\(Cri\)\s+\d+",
        ]
        citations = set()
        for p in citation_patterns:
            for m in re.finditer(p, text):
                citations.add(m.group(0))
        doc.citations_found = sorted(citations)[:20]

        # Dates
        date_patterns = [
            r"\d{1,2}[./-]\d{1,2}[./-]\d{2,4}",
            r"\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December),?\s+\d{4}",
        ]
        dates = set()
        for p in date_patterns:
            for m in re.finditer(p, text, re.IGNORECASE):
                dates.add(m.group(0))
        doc.dates_found = sorted(dates)[:20]

        # Parties (basic — look for vs/versus patterns)
        party_pattern = r"([\w\s\.]+?)\s+(?:vs?\.?|versus|v/s)\s+([\w\s\.]+?)(?:\s+on\s|\s*$|\s+\d)"
        for m in re.finditer(party_pattern, text[:2000], re.IGNORECASE):
            p1 = m.group(1).strip()
            p2 = m.group(2).strip()
            if len(p1) > 3 and len(p1) < 100:
                doc.parties_mentioned.append(p1)
            if len(p2) > 3 and len(p2) < 100:
                doc.parties_mentioned.append(p2)
        doc.parties_mentioned = doc.parties_mentioned[:10]

    def _guess_document_type(self, doc: ExtractedDocument) -> None:
        """Guess what type of legal document this is."""
        text = doc.text[:3000].lower()

        type_signals = [
            ("fir", ["first information report", "f.i.r", "fir no", "police station", "cognizable offence"]),
            ("charge_sheet", ["charge sheet", "chargesheet", "charge-sheet", "final report"]),
            ("court_order", ["it is ordered", "order dated", "this court orders", "hon'ble court"]),
            ("judgment", ["judgment", "judgement", "we hold", "appeal is", "petition is", "it is held"]),
            ("bail_application", ["bail application", "application for bail", "regular bail", "anticipatory bail"]),
            ("petition", ["writ petition", "criminal petition", "civil petition", "petition under"]),
            ("affidavit", ["affidavit", "i solemnly affirm", "sworn statement"]),
            ("complaint", ["complaint", "complainant", "prayer for"]),
            ("notice", ["legal notice", "notice under", "cease and desist"]),
            ("written_statement", ["written statement", "written_statement", "defence statement"]),
            ("appeal", ["appeal", "memorandum of appeal", "grounds of appeal"]),
            ("evidence_photo", []),  # Handled separately
        ]

        best_type = ""
        best_score = 0
        for doc_type, keywords in type_signals:
            score = sum(1 for kw in keywords if kw in text)
            if score > best_score:
                best_score = score
                best_type = doc_type

        doc.doc_type_guess = best_type or "legal_document"


# ── Format for Agent Context ─────────────────────────────────────────────

def format_documents_for_llm(docs: List[ExtractedDocument], max_total_chars: int = 15000) -> str:
    """
    Format multiple documents into a single context string for the LLM.
    
    Prioritizes: evidence descriptions, then shorter docs first,
    truncates long ones to fit within token budget.
    """
    if not docs:
        return ""

    # Sort: evidence photos first, then by length (shorter first)
    sorted_docs = sorted(docs, key=lambda d: (
        0 if d.evidence_description else 1,
        len(d.text),
    ))

    parts = ["=== UPLOADED CASE DOCUMENTS ===\n"]
    chars_used = len(parts[0])

    for i, doc in enumerate(sorted_docs):
        header = f"\n--- Document {i+1}: {doc.filename} ({doc.doc_type_guess or doc.file_type}) ---\n"
        chars_used += len(header)

        if chars_used >= max_total_chars:
            parts.append(f"\n[{len(sorted_docs) - i} more documents truncated for brevity]")
            break

        parts.append(header)

        if doc.evidence_description:
            desc = f"[Evidence Description: {doc.evidence_description}]\n"
            parts.append(desc)
            chars_used += len(desc)

        # Budget remaining chars for this doc's text
        remaining = max_total_chars - chars_used
        per_doc_budget = max(500, remaining // max(1, len(sorted_docs) - i))
        text_portion = doc.text[:per_doc_budget]
        if len(doc.text) > per_doc_budget:
            text_portion += "\n[...document truncated...]"

        parts.append(text_portion)
        chars_used += len(text_portion)

        # Add extracted entities
        if doc.sections_mentioned:
            entity_line = f"\n[Sections: {', '.join(doc.sections_mentioned[:10])}]"
            parts.append(entity_line)
            chars_used += len(entity_line)

    return "\n".join(parts)
